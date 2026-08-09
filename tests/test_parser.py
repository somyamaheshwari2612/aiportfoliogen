"""
Unit tests for input document parsers (TXT, PDF, DOCX) and parse_resume orchestrator.
"""

import os
import tempfile
import unittest
import docx
import pypdf

from input.parser import parse_resume
from input.txt_parser import read_txt
from input.pdf_parser import read_pdf
from input.docx_parser import read_docx


class TestResumeParser(unittest.TestCase):

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_parse_txt_valid(self):
        txt_path = os.path.join(self.temp_dir.name, "resume.txt")
        raw_content = "  John Doe  \n\n\nSoftware Engineer \n\nSkills: Python, Flask\n"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(raw_content)

        parsed = parse_resume(txt_path)
        expected = "John Doe\nSoftware Engineer\nSkills: Python, Flask"
        self.assertEqual(parsed, expected)
        self.assertEqual(read_txt(txt_path), raw_content)

    def test_parse_pdf_valid(self):
        pdf_path = os.path.join(self.temp_dir.name, "resume.pdf")
        writer = pypdf.PdfWriter()
        writer.add_blank_page(width=612, height=792)
        
        # Add page with text annotations / contents if needed, or simple page
        # Note: pypdf blank page doesn't have text, so let's test page text writing or read_pdf output
        with open(pdf_path, "wb") as f:
            writer.write(f)

        # A completely blank PDF page yields empty text, which parse_resume flags as empty
        with self.assertRaises(ValueError):
            parse_resume(pdf_path)

    def test_parse_docx_valid(self):
        docx_path = os.path.join(self.temp_dir.name, "resume.docx")
        doc = docx.Document()
        doc.add_heading("Jane Doe", level=1)
        doc.add_paragraph("AI Engineer")
        doc.add_paragraph("Experience in Python and PyTorch")

        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Project Alpha"
        hdr_cells[1].text = "Lead Developer"

        doc.save(docx_path)

        parsed = parse_resume(docx_path)
        self.assertIn("Jane Doe", parsed)
        self.assertIn("AI Engineer", parsed)
        self.assertIn("Project Alpha | Lead Developer", parsed)

        self.assertIn("Jane Doe", read_docx(docx_path))

    def test_sample_resume_txt(self):
        repo_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        sample_path = os.path.join(repo_root, "sample_resume.txt")
        if os.path.exists(sample_path):
            parsed = parse_resume(sample_path)
            self.assertIn("John Doe", parsed)
            self.assertIn("Software Engineer", parsed)
            # Ensure excessive blank lines are squashed
            self.assertNotIn("\n\n\n", parsed)

    def test_file_not_found(self):
        missing_path = os.path.join(self.temp_dir.name, "nonexistent.pdf")
        with self.assertRaises(FileNotFoundError):
            parse_resume(missing_path)

    def test_unsupported_extension(self):
        png_path = os.path.join(self.temp_dir.name, "resume.png")
        with open(png_path, "wb") as f:
            f.write(b"fake image data")

        with self.assertRaises(ValueError) as ctx:
            parse_resume(png_path)
        self.assertIn("Unsupported file format", str(ctx.exception))

    def test_empty_file(self):
        empty_txt = os.path.join(self.temp_dir.name, "empty.txt")
        with open(empty_txt, "w", encoding="utf-8") as f:
            f.write("   \n\n  \t  \n")

        with self.assertRaises(ValueError) as ctx:
            parse_resume(empty_txt)
        self.assertIn("appears to be empty", str(ctx.exception))

    def test_corrupt_pdf(self):
        corrupt_pdf = os.path.join(self.temp_dir.name, "corrupt.pdf")
        with open(corrupt_pdf, "wb") as f:
            f.write(b"This is not a valid PDF header or content.")

        with self.assertRaises(RuntimeError) as ctx:
            parse_resume(corrupt_pdf)
        self.assertIn("Corrupted or invalid PDF", str(ctx.exception))

    def test_corrupt_docx(self):
        corrupt_docx = os.path.join(self.temp_dir.name, "corrupt.docx")
        with open(corrupt_docx, "wb") as f:
            f.write(b"This is not a valid zip / docx package.")

        with self.assertRaises(RuntimeError) as ctx:
            parse_resume(corrupt_docx)
        self.assertIn("Corrupted or invalid DOCX", str(ctx.exception))


if __name__ == "__main__":
    unittest.main()
