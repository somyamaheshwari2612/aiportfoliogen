"""
Module for extracting text content from Microsoft Word (.docx) files using python-docx.
"""

import os
import docx
from docx.opc.exceptions import OpcError


def read_docx(path: str) -> str:
    """Extracts text content from paragraphs and tables in a DOCX document.

    Args:
        path (str): Path to the DOCX file.

    Returns:
        str: Extracted text content from the document.

    Raises:
        FileNotFoundError: If the DOCX file does not exist at path.
        RuntimeError: If the file is corrupted, invalid format, or unreadable.
    """
    if not os.path.exists(path) or not os.path.isfile(path):
        raise FileNotFoundError(f"DOCX file not found: {path}")

    try:
        doc = docx.Document(path)
        extracted_lines = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                extracted_lines.append(paragraph.text)

        # Extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    extracted_lines.append(" | ".join(row_text))

        return "\n".join(extracted_lines)
    except (OpcError, Exception) as err:
        raise RuntimeError(f"Corrupted or invalid DOCX file '{path}': {err}") from err



